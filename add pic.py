from docx.shared import Cm
from docx import Document
from PIL import Image, ExifTags


docAddPicName = input('input documnet name: ')
picNumbers = input('numbers of pictures: ')
docAddPic = Document(docAddPicName+'.docx')
paragraph = docAddPic.add_paragraph()

for i in range(int(picNumbers)):
    picName = input('name of pictures: ')
    try:
        image=Image.open(picName+'.jpg')

        for orientation in ExifTags.TAGS.keys():
            if ExifTags.TAGS[orientation]=='Orientation':
                break
        
        exif = image._getexif()

        try:
            if exif[orientation] == 3:
                image=image.rotate(180, expand=True)
            elif exif[orientation] == 6:
                image=image.rotate(270, expand=True)
            elif exif[orientation] == 8:
                image=image.rotate(90, expand=True)
            print('process done!')
        except:
            print('no need to process!')

        image.save(picName+'.jpg')
        image.close()

        run = paragraph.add_run()
        run.add_picture(picName+'.jpg', width=Cm(4))
        docAddPic.save('demo.docx')
        print('photo added!')

    except (AttributeError, KeyError, IndexError):
        # cases: image don't have getexif
        print('no need to process!')
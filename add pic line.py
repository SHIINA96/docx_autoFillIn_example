from docx.shared import Cm
from docx import Document

document = Document('2.docx')
paragraph = document.add_paragraph()
run = paragraph.add_run()
run.add_picture('2.jpg', width=Cm(6))
run_2 = paragraph.add_run()
run_2.add_picture('1.jpg', width=Cm(6))
run_3 = paragraph.add_run()
run_3.add_picture('3.jpg', width=Cm(6))
document.save('demo.docx')
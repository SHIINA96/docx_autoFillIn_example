# 自动添加多个照片到文件中

from docx.shared import Cm
from docx import Document
from PIL import Image, ExifTags

import glob, os

# 自动获取当前路径
dir_path = os.path.dirname(os.path.realpath(__file__))
# 创建图片文件名列表
picList = []


# 从路径中查找 .jpg 文件
os.chdir(dir_path)
for file in glob.glob("*.PNG"):
    i=0
    picList.insert(i, file)
    print(file)
    i = i+1
# print(picList)


docAddPicName = input('input documnet name: ')
docAddPic = Document(docAddPicName+'.docx')
paragraph = docAddPic.add_paragraph()

for picName in picList:
    # 检测图片是否需要修改角度信息
    try:
        image=Image.open(picName)
        for orientation in ExifTags.TAGS.keys():
            if ExifTags.TAGS[orientation]=='Orientation':
                break
        exif = image._getexif()
        try:
            if exif[orientation] == 3:
                image=image.rotate(180, expand=True)
            elif exif[orientation] == 6:
                image=image.rotate(270, expand=True)
            elif exif[orientation] == 8:
                image=image.rotate(90, expand=True)
            # print('process done!')
        except:
            # print('no need to process!')
            pass
        # 保存修改后的图片
        image.save(picName)
        image.close()
        # 添加图片到文档中
        run = paragraph.add_run()
        run.add_picture(picName, width=Cm(10))
        docAddPic.save('demo.docx')
        # print('photo added!')

    except (AttributeError, KeyError, IndexError):
        # cases: image don't have getexif
        print('no need to process!')
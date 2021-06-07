from docxtpl import DocxTemplate
from docx import Document
from docx2pdf import convert
from docx.shared import Cm
from PIL import Image, ExifTags

import datetime
import glob, os

# 获取当前时间
today = datetime.date.today()
# 自动获取当前路径
dir_path = os.path.dirname(os.path.realpath(__file__))
# 创建图片文件名列表
picList = []


# 从路径中查找 .jpg 文件
os.chdir(dir_path)
for file in glob.glob("*.jpg"):
    i=0
    picList.insert(i, file)
    # print(file)
    i = i+1
# print(picList)

# 从照片中获取日期
try:
    pict = Image.open(picList[0])
    exif_data = pict._getexif()
    picDate = exif_data[36867]
    pict.close()
except:
    pict = Image.open(picList[1])
    exif_data = pict._getexif()
    picDate = exif_data[36867]
    pict.close()

name = input('输入文件名：')
student_name = input('输入学生名字：')
communication = int(input('思辨与交流能力: '))
creation = int(input('反思与创新能力：'))
co_operation = int(input('合作与互助能力：'))
solvability = int(input('问题解决能力：'))
thoughts = int(input('编程思维：'))
lecturer_comments = input('输入教师评价：')
year = picDate[0:4]
month = picDate[5:7]
day = picDate[8:10]


doc = DocxTemplate(name+'.docx') #加载模板文件
document = Document(name+'.docx')

data_dic = {
    'student_name' : student_name,
    'communication' : '★'*(communication-1),
    'creation' : '★'*(creation-1),
    'co_operation' : '★'*(co_operation-1),
    'solvability' : '★'*(solvability-1),
    'thoughts' : '★'*(thoughts-1),
    'lecturer_comments' : lecturer_comments,
    'year' : year,
    'month' : month,
    'day' : day
}
doc.render(data_dic) #填充数据
table = doc.tables[0] # 获取课程名称
lessonName = table.cell(1, 1).text
documentName = student_name+lessonName+str(year)+str(month)+str(day)+'课评报告.docx'
doc.save(documentName) #保存目标文件


docAddPicName = documentName
docAddPic = Document(docAddPicName)
paragraph = docAddPic.add_paragraph()
# 检测图片是否需要旋转
n = 0
for picName in picList:
    try:
        image=Image.open(picName)
        for orientation in ExifTags.TAGS.keys():
            if ExifTags.TAGS[orientation]=='Orientation':
                break
        exif = image._getexif()
        try:
            if exif[orientation] == 3:
                image=image.rotate(180, expand=True)
            elif exif[orientation] == 6:
                image=image.rotate(270, expand=True)
            elif exif[orientation] == 8:
                image=image.rotate(90, expand=True)
            # print('process done!')
        except:
            # print('no need to process!')
            pass
        image.save(picName)
        image.close()

        # 添加照片到文件中
        run = paragraph.add_run()
        run.add_picture(picName, width=Cm(6))
        docAddPic.save(documentName)
        print('photo added!')
        image.close()

        # 更改照片名
        picNameStandard = year+month+day+'_'+str(n)+'.jpg'
        os.rename(picName, picNameStandard)
        n = n+1

    except (AttributeError, KeyError, IndexError):
        # cases: image don't have getexif
        print('no need to process!')


# 生成PDF
generatePDF = input('是否生成PDF？ YES｜NO\n')
if generatePDF == 'YES' or 'yes':
    convert(documentName)
else:
    pass